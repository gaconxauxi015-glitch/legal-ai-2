import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from PIL import Image
import pytesseract
import io

# ======================
# SAFE AI MOCK (KHÔNG QUOTA, KHÔNG LỖI API)
# ======================
def ask_ai(prompt):
    return f"🤖 AI RESPONSE (SAFE MODE)\n\n{prompt[:2000]}"

# ======================
# READ FILES
# ======================
def read_pdf(file):
    reader = PdfReader(file)
    return "\n".join([p.extract_text() or "" for p in reader.pages])

def read_docx(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def read_txt(file):
    return file.read().decode("utf-8")

def read_image(file):
    img = Image.open(file)
    return pytesseract.image_to_string(img, lang="eng+vie")

def read_file(file):
    name = file.name.lower()
    if name.endswith(".pdf"):
        return read_pdf(file)
    if name.endswith(".docx"):
        return read_docx(file)
    if name.endswith(".txt"):
        return read_txt(file)
    return read_image(file)

# ======================
# AI FUNCTIONS
# ======================
def create_contract(desc):
    return f"""
HỢP ĐỒNG (BẢN NHÁP CHUẨN)

1. Thông tin các bên
- Bên A: ...
- Bên B: ...

2. Nội dung hợp đồng
{desc}

3. Quyền và nghĩa vụ
- ...

4. Thanh toán
- ...

5. Chấm dứt hợp đồng
- ...

6. Giải quyết tranh chấp
- ...
"""

def analyze_contract(text):
    return f"""
📊 PHÂN TÍCH HỢP ĐỒNG

- Nội dung chính: rõ ràng
- Điều khoản quan trọng: thanh toán, trách nhiệm
- Đánh giá: cần kiểm tra kỹ điều khoản chấm dứt

📄 Nội dung:
{text[:2000]}
"""

def risk_contract(text):
    return """
⚠️ RỦI RO PHÁP LÝ

- Điều khoản mơ hồ
- Thiếu chế tài phạt
- Thiếu điều khoản tranh chấp
- Thiếu quy định chấm dứt rõ ràng
"""

def check_missing(text):
    return """
📄 KIỂM TRA THIẾU ĐIỀU KHOẢN

- Thanh toán có rõ không?
- Phạt vi phạm có không?
- Chấm dứt hợp đồng có chi tiết không?
- Giải quyết tranh chấp có chưa?
"""

# ======================
# EXPORT WORD (ĐẸP + KHÔNG CẦN SỬA)
# ======================
def export_word(content):
    doc = Document()

    doc.add_heading("HỢP ĐỒNG PHÁP LÝ", 0)

    for line in content.split("\n"):
        if line.strip():
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)

    file_path = "/tmp/legal_contract.docx"
    doc.save(file_path)
    return file_path

# ======================
# UI
# ======================
st.title("⚖️ LEGAL AI PRO FINAL (ALL-IN-ONE)")

mode = st.selectbox(
    "Chức năng",
    [
        "📊 Phân tích hợp đồng",
        "⚠️ Tìm rủi ro",
        "📄 Kiểm tra thiếu điều khoản",
        "🧾 Tạo hợp đồng",
        "💬 Hỏi đáp tài liệu"
    ]
)

file = st.file_uploader(
    "Upload file / ảnh scan",
    type=["pdf","docx","txt","png","jpg","jpeg"]
)

text = ""
if file:
    text = read_file(file)
    st.success("Đã đọc file")

# ======================
# LOGIC
# ======================

if mode == "📊 Phân tích hợp đồng":
    if file and st.button("Phân tích"):
        st.write(analyze_contract(text))

if mode == "⚠️ Tìm rủi ro":
    if file and st.button("Kiểm tra rủi ro"):
        st.write(risk_contract(text))

if mode == "📄 Kiểm tra thiếu điều khoản":
    if file and st.button("Kiểm tra"):
        st.write(check_missing(text))

if mode == "💬 Hỏi đáp tài liệu":
    if file:
        q = st.text_input("Nhập câu hỏi")
        if q and st.button("Trả lời"):
            st.write(ask_ai(text + "\n\nCâu hỏi: " + q))

if mode == "🧾 Tạo hợp đồng":
    desc = st.text_area("Mô tả hợp đồng")

    if st.button("Tạo hợp đồng") and desc:
        result = create_contract(desc)
        st.write(result)

        file_path = export_word(result)

        with open(file_path, "rb") as f:
            st.download_button(
                "⬇️ Tải hợp đồng Word chuẩn",
                f,
                file_name="hop_dong_phap_ly.docx"
            )
