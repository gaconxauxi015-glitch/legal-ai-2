from docx import Document
from docx.shared import Pt

def export_word(content):

    doc = Document()

    doc.add_heading("HỢP ĐỒNG", 0)

    for line in content.split("\\n"):

        if line.strip():

            p = doc.add_paragraph(line)

            if p.runs:
                p.runs[0].font.size = Pt(11)

    file_path = "/tmp/hop_dong.docx"

    doc.save(file_path)

    return file_path
